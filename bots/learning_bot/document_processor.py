import logging
from typing import List, Optional
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles loading and processing of documents for the RAG system."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks for splitting
            chunk_overlap: Overlap between chunks
        """
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def load_document(self, file_path: str) -> Optional[str]:
        """
        Load a document from file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text from the document or None if loading fails
        """
        try:
            file_path = Path(file_path)
            if file_path.suffix.lower() == '.pdf':
                return self._load_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._load_docx(file_path)
            else:
                logger.error(f"Unsupported file format: {file_path.suffix}")
                return None
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return None
    
    def _load_pdf(self, file_path: Path) -> str:
        """Load text from a PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf = PdfReader(file)
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _load_docx(self, file_path: Path) -> str:
        """Load text from a DOCX file."""
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def process_document(self, text: str) -> List[LangchainDocument]:
        """
        Process document text into chunks for the RAG system.
        
        Args:
            text: Document text to process
            
        Returns:
            List of LangchainDocument chunks
        """
        try:
            chunks = self.text_splitter.split_text(text)
            return [
                LangchainDocument(
                    page_content=chunk,
                    metadata={"source": "document"}
                )
                for chunk in chunks
            ]
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            return [] 